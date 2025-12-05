#!/usr/bin/env python3
"""
Convert CanOEs User Manual from Markdown to PDF and DOCX formats.

Usage:
    python convert_manual.py
"""

import os
import markdown
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(SCRIPT_DIR, "CanOEs_User_Manual.md")
PDF_FILE = os.path.join(SCRIPT_DIR, "CanOEs_User_Manual.pdf")
DOCX_FILE = os.path.join(SCRIPT_DIR, "CanOEs_User_Manual.docx")

# CSS for PDF styling
PDF_CSS = """
@page {
    size: A4;
    margin: 2cm;
}

body {
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #333;
}

h1 {
    color: #1a5490;
    font-size: 24pt;
    border-bottom: 2px solid #1a5490;
    padding-bottom: 10px;
    margin-top: 30px;
}

h2 {
    color: #2c7bb6;
    font-size: 18pt;
    margin-top: 25px;
    border-bottom: 1px solid #ddd;
    padding-bottom: 5px;
}

h3 {
    color: #4a9fd4;
    font-size: 14pt;
    margin-top: 20px;
}

h4 {
    color: #666;
    font-size: 12pt;
    margin-top: 15px;
}

table {
    border-collapse: collapse;
    width: 100%;
    margin: 15px 0;
}

th, td {
    border: 1px solid #ddd;
    padding: 8px 12px;
    text-align: left;
}

th {
    background-color: #1a5490;
    color: white;
}

tr:nth-child(even) {
    background-color: #f9f9f9;
}

code {
    background-color: #f4f4f4;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 10pt;
}

pre {
    background-color: #f4f4f4;
    padding: 15px;
    border-radius: 5px;
    overflow-x: auto;
    font-family: 'Consolas', 'Courier New', monospace;
    font-size: 9pt;
    line-height: 1.4;
}

blockquote {
    border-left: 4px solid #1a5490;
    margin: 15px 0;
    padding: 10px 20px;
    background-color: #f9f9f9;
}

hr {
    border: none;
    border-top: 1px solid #ddd;
    margin: 20px 0;
}

ul, ol {
    margin: 10px 0;
    padding-left: 25px;
}

li {
    margin: 5px 0;
}
"""


def convert_to_pdf():
    """Convert Markdown to PDF using WeasyPrint."""
    print("Converting to PDF...")
    
    # Read markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML
    html_content = markdown.markdown(
        md_content, 
        extensions=['tables', 'fenced_code', 'toc']
    )
    
    # Wrap in HTML document
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>CanOEs User Manual</title>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Generate PDF
    HTML(string=full_html).write_pdf(PDF_FILE, stylesheets=[CSS(string=PDF_CSS)])
    print(f"  Created: {PDF_FILE}")


def convert_to_docx():
    """Convert Markdown to DOCX using python-docx."""
    print("Converting to DOCX...")
    
    # Read markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Parse markdown line by line
    lines = md_content.split('\n')
    in_code_block = False
    in_table = False
    table_data = []
    
    for line in lines:
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style.font.name = 'Consolas'
            p.style.font.size = Pt(9)
            continue
        
        # Tables
        if '|' in line and not line.startswith('```'):
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                table_data.append(cells)
            continue
        elif table_data:
            # End of table, render it
            if len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell_text
            table_data = []
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        elif line.strip():
            # Regular paragraph - handle inline formatting
            p = doc.add_paragraph()
            # Simple approach - just add text
            text = line.replace('**', '').replace('*', '').replace('`', '')
            p.add_run(text)
    
    # Handle remaining table
    if table_data:
        if len(table_data) > 0:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_text
    
    # Save
    doc.save(DOCX_FILE)
    print(f"  Created: {DOCX_FILE}")


def main():
    print("=" * 50)
    print("  CanOEs User Manual Converter")
    print("=" * 50)
    print()
    
    if not os.path.exists(MD_FILE):
        print(f"Error: {MD_FILE} not found!")
        return
    
    try:
        convert_to_pdf()
    except Exception as e:
        print(f"  PDF Error: {e}")
    
    try:
        convert_to_docx()
    except Exception as e:
        print(f"  DOCX Error: {e}")
    
    print()
    print("Done!")


if __name__ == "__main__":
    main()
